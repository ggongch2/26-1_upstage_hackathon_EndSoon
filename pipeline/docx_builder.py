"""Reassemble translated elements into a Word document preserving order/layout."""
from __future__ import annotations

import base64
import logging
import re
from io import BytesIO
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt

from .translator import TranslatedElement

log = logging.getLogger(__name__)

HEADING_LEVELS = {
    "heading1": 1,
    "heading2": 2,
    "heading3": 3,
    "title": 0,
}

IMAGE_CATEGORIES = {"figure", "chart"}
TABLE_CATEGORIES = {"table"}


def _add_image_from_base64(doc: Document, b64: str, max_width_inches: float = 5.5) -> None:
    if not b64:
        return
    data = b64
    if "," in data:
        data = data.split(",", 1)[1]
    try:
        raw = base64.b64decode(data)
    except Exception:
        log.warning("invalid base64 image; skipping")
        return
    bio = BytesIO(raw)
    try:
        doc.add_picture(bio, width=Inches(max_width_inches))
    except Exception:
        log.exception("add_picture failed; skipping image")


def _add_html_table(doc: Document, html: str) -> None:
    if not html.strip():
        return
    soup = BeautifulSoup(html, "lxml")
    table_tag = soup.find("table")
    if not table_tag:
        doc.add_paragraph(soup.get_text("\n", strip=True))
        return

    rows = table_tag.find_all("tr")
    if not rows:
        return
    cols = max(len(r.find_all(["td", "th"])) for r in rows)
    if cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        for ci in range(cols):
            text = cells[ci].get_text(" ", strip=True) if ci < len(cells) else ""
            table.cell(ri, ci).text = text


_LATEX_INLINE = re.compile(r"(\$[^$\n]+\$|\\\([^)]+\\\))")


def _add_equation_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)


def build_docx(
    translated: list[TranslatedElement],
    output_path: str | Path,
    *,
    title: str | None = None,
) -> Path:
    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    last_page = None
    for tr in translated:
        elem = tr.element
        if last_page is not None and elem.page != last_page:
            doc.add_paragraph()
        last_page = elem.page

        cat = elem.category.lower()
        if cat in IMAGE_CATEGORIES:
            _add_image_from_base64(doc, elem.base64 or "")
            if tr.translated_text.strip():
                cap = doc.add_paragraph(tr.translated_text.strip())
                cap.italic = True
            continue

        if cat in TABLE_CATEGORIES:
            _add_html_table(doc, tr.translated_html or elem.html)
            continue

        if cat == "equation":
            _add_equation_paragraph(doc, elem.text or elem.markdown)
            continue

        if cat in HEADING_LEVELS:
            doc.add_heading(tr.translated_text or elem.text, level=HEADING_LEVELS[cat])
            continue

        if cat == "list":
            for line in (tr.translated_text or "").splitlines():
                line = line.strip()
                if not line:
                    continue
                doc.add_paragraph(line.lstrip("-•* "), style="List Bullet")
            continue

        text = tr.translated_text.strip()
        if not text:
            continue
        doc.add_paragraph(text)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def docx_to_pdf(docx_path: str | Path, pdf_path: str | Path | None = None) -> Path | None:
    """Optional PDF export — requires MS Word installed (docx2pdf, Windows/Mac)."""
    try:
        from docx2pdf import convert
    except ImportError:
        log.warning("docx2pdf not available; skipping PDF export")
        return None
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path) if pdf_path else docx_path.with_suffix(".pdf")
    try:
        convert(str(docx_path), str(pdf_path))
    except Exception:
        log.exception("docx2pdf conversion failed")
        return None
    return pdf_path
